from pathlib import Path
from zipfile import ZipFile
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

OUT = Path('output')
OUT.mkdir(exist_ok=True)
DOCX = OUT / '20260715我们的秘密基地——东方社区23幢架空层粉刷儿童议事会会议记录表.docx'
PACKAGE = Path('package')


def set_cell_shading(cell, fill='F2F2F2'):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn('w:shd'))
    if shd is None:
        shd = OxmlElement('w:shd')
        tc_pr.append(shd)
    shd.set(qn('w:fill'), fill)


def set_cell_text(cell, text, bold=False, size=12, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(text)
    r.bold = bold
    r.font.name = '宋体'
    r._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    r.font.size = Pt(size)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_body(doc, text, bold_prefix=None):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.line_spacing = 1.6
    p.paragraph_format.space_after = Pt(4)
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        r1.bold = True
        r1.font.name = '宋体'
        r1._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        r1.font.size = Pt(14)
        text = text[len(bold_prefix):]
    r = p.add_run(text)
    r.font.name = '宋体'
    r._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    r.font.size = Pt(14)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    r.bold = True
    r.font.name = '黑体'
    r._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    r.font.size = Pt(15 if level == 1 else 14)
    return p


doc = Document()
section = doc.sections[0]
section.page_width = Cm(21)
section.page_height = Cm(29.7)
section.top_margin = Cm(2.2)
section.bottom_margin = Cm(2.2)
section.left_margin = Cm(2.0)
section.right_margin = Cm(2.0)

styles = doc.styles
styles['Normal'].font.name = '宋体'
styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
styles['Normal'].font.size = Pt(14)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(18)
r = p.add_run('“我们的秘密基地”——东方社区23幢架空层粉刷儿童议事会\n会议记录表')
r.font.name = '方正小标宋简体'
r._element.rPr.rFonts.set(qn('w:eastAsia'), '方正小标宋简体')
r.font.size = Pt(22)

info = doc.add_table(rows=3, cols=4)
info.alignment = WD_TABLE_ALIGNMENT.CENTER
info.style = 'Table Grid'
widths = [Cm(2.2), Cm(6.1), Cm(2.2), Cm(6.1)]
for row in info.rows:
    for idx, cell in enumerate(row.cells):
        cell.width = widths[idx]
labels = [
    ('会议主题', '“我们的秘密基地”——东方社区23幢架空层粉刷儿童议事会', '会议时间', '2026年7月15日14:00—15:00'),
    ('会议地点', '东方社区活动室', '主持人', '社区社工'),
    ('参会人员', '东方社区儿童、家长志愿者、社区社工室工作人员', '参会人数', '以签到表为准'),
]
for i, row_data in enumerate(labels):
    for j, value in enumerate(row_data):
        is_label = j in (0, 2)
        set_cell_text(info.cell(i, j), value, bold=is_label, size=11 if not is_label else 12,
                      align=WD_ALIGN_PARAGRAPH.CENTER if is_label else WD_ALIGN_PARAGRAPH.LEFT)
        if is_label:
            set_cell_shading(info.cell(i, j))

doc.add_paragraph()

add_heading(doc, '一、开场导入：认识共同的“小区客厅”')
add_body(doc, '活动开始，社工以“我们的秘密基地”为主题欢迎小朋友参加议事会，并通过图片对比和生活化提问，引导大家认识架空层。社工介绍，架空层既是居民日常经过、休息和交流的公共空间，也是需要保持消防通道畅通、环境整洁和秩序安全的“小区客厅”。')
add_body(doc, '社工重点说明，东方社区将对23幢架空层进行墙面粉刷和环境改善。本次活动不仅是让小朋友了解粉刷工作的意义，更重要的是邀请大家共同商量：粉刷完成后，我们怎样保护墙面、爱护设施、维护卫生，让架空层长期保持安全、整洁和美观。')
add_body(doc, '同时，社工明确指出，后续维护将以23幢架空层为重点，但爱护公共空间不能只局限于一栋楼。社区内其他楼栋的架空层同样属于居民共同使用的公共区域，也需要大家在日常生活中共同维护。')

add_heading(doc, '二、主题学习：为什么要爱护架空层')
add_body(doc, '社工结合PPT内容，从安全、健康、环境和邻里感受四个方面进行讲解。电动车乱停乱放、杂物占用通道会影响正常通行，也可能堵塞消防通道；在架空层及楼道内违规充电、堆放易燃杂物，还可能增加火灾风险。保持通道畅通，就是守护居民共同的安全。')
add_body(doc, '地面垃圾、污渍和长期堆放的杂物不仅影响美观，也容易产生异味和卫生隐患。干净、通风、有序的架空层，能够让居民进出更加方便，也能让邻里在公共空间中感到舒适和安心。')
add_body(doc, '社工提醒小朋友，粉刷后的墙面、公共座椅、栏杆及其他设施属于全体居民共同所有，不能随意涂画、踢踹或者损坏；在架空层内也不应追逐打闹、开展危险游戏，要在保护自己的同时避免影响他人。')

add_heading(doc, '三、儿童议事：我们可以为架空层做些什么')
add_body(doc, '在讨论环节，社工围绕“自己可以做到什么”“怎样帮助维护23幢架空层”“怎样带动周边的人一起参与”三个问题，引导小朋友逐一发表意见。小朋友们结合日常观察，讨论重点集中在爱护环境、完成力所能及的小事以及带动家人和伙伴共同维护环境秩序。')

add_heading(doc, '（一）从自己做起，养成爱护环境的习惯', level=2)
add_body(doc, '小朋友们认为，爱护架空层首先要管好自己的行为，做到不乱扔果皮纸屑、不随地吐痰，不在新粉刷的墙面上乱涂乱画，不踢踹座椅和栏杆，不破坏公共设施。同时，不在架空层及消防通道内追逐打闹，不在车辆进出区域进行危险游戏。')
add_body(doc, '有小朋友表示，新粉刷的墙面要像家里的墙面一样爱护，“不能因为是公共地方就随便弄脏”。大家一致认为，只有每个人先从自己做起，23幢架空层的改善成果才能保持得更久，其他楼栋的架空层也才能一直保持整洁。')

add_heading(doc, '（二）主动完成力所能及的小事', level=2)
add_body(doc, '小朋友们提出，看到地面上的纸屑、塑料袋等普通垃圾，可以在确保安全的情况下主动捡起并放入垃圾桶；参加社区志愿活动时，可以在成年人带领下擦拭公共座椅、整理活动用品、清理小广告，做自己能够完成的事情。')
add_body(doc, '对于电动车堵塞通道、大件杂物堆放、公共设施损坏等问题，小朋友们认为不应自行搬动或处理，而应及时告诉家长、社区工作人员或物业人员。社工强调，儿童参与公共环境维护要以安全为前提，既要主动，也要学会正确求助。')

add_heading(doc, '（三）礼貌提醒他人，带动周边居民参与', level=2)
add_body(doc, '小朋友们认为，爱护环境不仅是“不做不文明的事”，还可以用友善的方式影响身边的人。看到家人或伙伴不小心乱扔垃圾时，可以礼貌提醒“垃圾要回家”；看到小伙伴在墙面上乱画、在架空层里进行危险游戏时，可以及时劝阻，并在需要时寻求成年人帮助。')
add_body(doc, '有小朋友提出，可以邀请爸爸妈妈、兄弟姐妹和朋友一起参加社区清洁活动，把今天学到的架空层爱护知识告诉身边的人。大家认为，虽然儿童的力量有限，但每一次主动捡拾、每一句礼貌提醒，都能让更多人关注公共环境，共同维护社区秩序。')

add_heading(doc, '（四）以23幢为重点，延伸维护其他架空层', level=2)
add_body(doc, '围绕后续维护范围，参会儿童形成一致意见：23幢是本次粉刷和环境改善的重点楼栋，后续要重点关注墙面保护、垃圾清理、公共设施使用和通道秩序，把23幢架空层维护成整洁有序的示范空间。')
add_body(doc, '与此同时，无论经过社区哪一栋楼的架空层，都应遵守相同的公共环境规则。大家不能只爱护23幢，也应关注其他楼栋架空层的卫生、安全和秩序，发现问题及时反馈，带动更多居民共同保护社区公共空间。')

add_heading(doc, '四、形成儿童爱护架空层行动约定')
add_body(doc, '经过讨论，社工将小朋友们的意见梳理为《架空层爱护小卫士行动约定》：')
actions = [
    '不乱扔垃圾，不随地吐痰，保持公共区域整洁；',
    '不在墙面和公共设施上乱涂乱画，不损坏公物；',
    '不在架空层内追逐打闹，不占用或堵塞消防通道；',
    '看到普通垃圾主动捡拾，参与力所能及的清洁行动；',
    '发现杂物堆放、电动车占道或设施损坏等问题，及时向成年人反馈；',
    '礼貌提醒家人和伙伴，邀请更多人共同维护环境秩序；',
    '以23幢架空层为重点，将爱护环境的行动延伸到社区其他楼栋架空层。',
]
for idx, item in enumerate(actions, 1):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.74)
    p.paragraph_format.first_line_indent = Cm(-0.74)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(f'{idx}. {item}')
    r.font.name = '宋体'
    r._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    r.font.size = Pt(14)
add_body(doc, '社工将本次儿童议事形成的共识概括为：“从我做起、力所能及、友善提醒、共同维护。”小朋友们共同表示愿意担任架空层爱护小卫士，用自己的实际行动守护社区环境。')

add_heading(doc, '五、创意手工：石膏冰箱贴涂色')
add_body(doc, '议事讨论结束后，社工带领小朋友开展石膏冰箱贴创意手工。社工介绍颜料、画笔和石膏白坯的使用方法，提醒大家注意桌面整洁，避免颜料沾到衣物，并鼓励小朋友自由选择颜色和图案，为冰箱贴绘制表情、花纹和装饰细节。')
add_body(doc, '制作过程中，小朋友们相互交流创作想法、分享颜料，在轻松的氛围中完成了各具特色的作品。社工引导大家把冰箱贴带回家，作为“爱护家园”的小小提醒，将活动中形成的环保约定分享给家人。')

add_heading(doc, '六、会议总结及后续安排')
add_body(doc, '活动最后，社工对本次议事成果进行总结，肯定小朋友能够从公共安全、环境卫生、设施保护和带动他人参与等角度提出具体建议。社工指出，粉刷能够让架空层焕然一新，但真正决定环境能否长期保持的，是每一位居民日常的使用习惯和持续行动。')
add_body(doc, '后续，东方社区将以23幢架空层为主要维护对象，持续关注墙面保护、公共设施使用、垃圾清理和通道秩序，并结合实际情况将相关维护经验延伸至其他楼栋架空层。社区也将继续鼓励儿童及家庭参与环境维护和志愿服务，让更多居民从“公共空间使用者”转变为“公共环境守护者”。')
add_body(doc, '全体参会人员手持完成的冰箱贴作品合影留念，本次会议顺利结束。')

doc.add_page_break()
add_heading(doc, '会议照片')
for _ in range(2):
    t = doc.add_table(rows=1, cols=1)
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = t.cell(0, 0)
    cell.height = Cm(6)
    set_cell_text(cell, '活动照片粘贴处', size=12, align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph()
add_heading(doc, '与会人员签名')
t = doc.add_table(rows=1, cols=1)
t.style = 'Table Grid'
t.alignment = WD_TABLE_ALIGNMENT.CENTER
cell = t.cell(0, 0)
cell.height = Cm(7)
set_cell_text(cell, '签到表粘贴处', size=12, align=WD_ALIGN_PARAGRAPH.CENTER)

doc.add_paragraph()
p = doc.add_paragraph('记录人：________________________')
p.paragraph_format.space_before = Pt(10)
p = doc.add_paragraph('记录日期：2026年7月15日')

doc.save(DOCX)
PACKAGE.mkdir(exist_ok=True)
with ZipFile(DOCX) as zf:
    zf.extractall(PACKAGE)
print(DOCX)
